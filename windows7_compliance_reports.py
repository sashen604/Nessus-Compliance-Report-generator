"""Generate Windows 7 Compliance Reports (Summary and Details) in Excel and Word formats."""
from __future__ import annotations

import argparse
import re
import xml.etree.ElementTree as ET
from typing import List, Optional, Tuple
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

CHECK_PATTERN = re.compile(r"^(?P<check>\S+)\s+\((?P<level>L[12])\)\s+(?P<desc>.+)$")
WINDOWS_7_PATTERN = re.compile(r"Windows\s+7|CIS_MS_Windows_7|win.*7", re.IGNORECASE)


def _strip_text(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    value = value.strip()
    return value if value else None


def _find_text(element: ET.Element, tag: str) -> Optional[str]:
    child = element.find(tag)
    if child is not None and child.text is not None:
        return _strip_text(child.text)

    for child in element:
        if child.tag.endswith(tag):
            if child.text is not None:
                return _strip_text(child.text)

    return None


def is_windows_7_check(check_name: Optional[str], audit_file: Optional[str]) -> bool:
    """Check if this is a Windows 7 compliance check."""
    if not check_name and not audit_file:
        return False
    
    if check_name and WINDOWS_7_PATTERN.search(check_name):
        return True
    
    if audit_file and "Windows_7" in audit_file:
        return True
    
    return False


def extract_windows7_items(scan_path: str) -> Tuple[Optional[str], List[dict]]:
    """Extract Windows 7 compliance items from Nessus scan."""
    tree = ET.parse(scan_path)
    root = tree.getroot()

    host_ip: Optional[str] = None
    for report_host in root.iter("ReportHost"):
        host_ip = report_host.attrib.get("name") or _find_text(report_host, "host-ip")
        if host_ip:
            break

    items: List[dict] = []

    for report_item in root.iter("ReportItem"):
        compliance_text = _find_text(report_item, "compliance")
        compliance_result = _find_text(report_item, "compliance-result")
        if compliance_text and compliance_text.lower() != "true" and not compliance_result:
            continue
        if compliance_text is None and compliance_result is None:
            continue

        check_name = _find_text(report_item, "compliance-check-name")
        audit_file = _find_text(report_item, "cm:compliance-audit-file")
        
        # Only process Windows 7 compliance checks
        if not is_windows_7_check(check_name, audit_file):
            continue

        level = _find_text(report_item, "cm:compliance-benchmark-profile")
        if not level:
            # Try to extract from check name
            match = re.search(r"Level_([12])", check_name or "")
            level = f"L{match.group(1)}" if match else "L2"

        info_text = _find_text(report_item, "compliance-info") or ""
        solution_text = _find_text(report_item, "compliance-solution") or ""
        impact_text = _find_text(report_item, "compliance-impact") or ""

        if not impact_text and solution_text:
            split_token = "Impact:"
            if split_token in solution_text:
                solution_part, impact_part = solution_text.split(split_token, 1)
                solution_text = solution_part.strip()
                impact_text = impact_part.strip()

        items.append(
            {
                "check_number": _find_text(report_item, "cm:compliance-check-id") or "",
                "level": level,
                "description": check_name or "",
                "result": compliance_result or _find_text(report_item, "cm:compliance-actual-value") or "Unknown",
                "info": info_text,
                "solution": solution_text,
                "impact": impact_text,
                "policy_value": _find_text(report_item, "cm:compliance-policy-value") or "",
                "actual_value": _find_text(report_item, "cm:compliance-actual-value") or "",
                "benchmark": _find_text(report_item, "cm:compliance-benchmark-name") or "CIS Windows 7",
            }
        )

    return host_ip, items


def count_l1_l2(items: List[dict]) -> Tuple[int, int, int]:
    """Count L1, L2, and total items."""
    l1 = sum(1 for item in items if item.get("level") == "L1")
    l2 = sum(1 for item in items if item.get("level") == "L2")
    return l1, l2, len(items)


def set_cell_bg(cell, color: str = "D9E1F2") -> None:
    """Set cell background color in DOCX."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def build_summary_word(items: List[dict], output_path: str, host_ip: Optional[str] = None) -> None:
    """Build Windows 7 summary report in Word format."""
    document = Document()
    
    if host_ip:
        p = document.add_paragraph(f"Host: {host_ip}")
        p.runs[0].bold = True
    
    document.add_heading("Windows 7 Compliance Summary Report", level=1)
    
    l1, l2, total = count_l1_l2(items)
    
    document.add_heading("Summary Statistics", level=2)
    document.add_paragraph(f"L1 Checks: {l1}")
    document.add_paragraph(f"L2 Checks: {l2}")
    document.add_paragraph(f"Total Checks: {total}")
    
    # Create summary table
    document.add_heading("Compliance Summary Table", level=2)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    
    headers = ["Check ID", "L1/L2", "Description", "Status", "Benchmark"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header)
        run.bold = True
        set_cell_bg(cell, "4472C4")
    
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.get("check_number", ""))[:20]
        row_cells[1].text = str(item.get("level", ""))
        row_cells[2].text = str(item.get("description", ""))[:50]
        row_cells[3].text = str(item.get("result", ""))
        row_cells[4].text = str(item.get("benchmark", ""))
    
    document.save(output_path)


def build_summary_excel(items: List[dict], output_path: str, host_ip: Optional[str] = None) -> None:
    """Build Windows 7 summary report in Excel format."""
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Windows 7 Summary"
    
    # Define styles
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Title section
    ws_title = workbook.create_sheet("Summary", 0)
    ws_title['A1'] = "Windows 7 Compliance Summary Report"
    ws_title['A1'].font = Font(bold=True, size=14)
    
    if host_ip:
        ws_title['A2'] = f"Host: {host_ip}"
    
    l1, l2, total = count_l1_l2(items)
    ws_title['A4'] = "Summary Statistics"
    ws_title['A4'].font = Font(bold=True)
    ws_title['A5'] = f"L1 Checks: {l1}"
    ws_title['A6'] = f"L2 Checks: {l2}"
    ws_title['A7'] = f"Total Checks: {total}"
    
    # Data sheet
    data_sheet = worksheet
    data_sheet.title = "Compliance Data"
    
    # Column widths
    column_widths = [15, 8, 40, 12, 20]
    for i, width in enumerate(column_widths, 1):
        data_sheet.column_dimensions[chr(64 + i)].width = width
    
    # Headers
    headers = ["Check ID", "Level", "Description", "Status", "Benchmark"]
    for col_num, header in enumerate(headers, 1):
        cell = data_sheet.cell(row=1, column=col_num)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Data rows
    for row_num, item in enumerate(items, 2):
        row_data = [
            str(item.get("check_number", ""))[:20],
            str(item.get("level", "")),
            str(item.get("description", "")),
            str(item.get("result", "")),
            str(item.get("benchmark", "")),
        ]
        
        for col_num, value in enumerate(row_data, 1):
            cell = data_sheet.cell(row=row_num, column=col_num)
            cell.value = value
            cell.border = border
            cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
    
    workbook.save(output_path)


def build_details_word(items: List[dict], output_path: str, host_ip: Optional[str] = None) -> None:
    """Build Windows 7 detailed report in Word format."""
    document = Document()
    
    if host_ip:
        p = document.add_paragraph(f"Host: {host_ip}")
        p.runs[0].bold = True
    
    document.add_heading("Windows 7 Compliance Detailed Report", level=1)
    
    l1, l2, total = count_l1_l2(items)
    document.add_paragraph(f"L1 Checks: {l1}")
    document.add_paragraph(f"L2 Checks: {l2}")
    document.add_paragraph(f"Total: {total}")
    
    for item in items:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        
        # Header row with check info
        header_cells = table.rows[0].cells
        header_cells[0].merge(header_cells[2])
        
        p = header_cells[0].paragraphs[0]
        p.clear()
        run = p.add_run(f"[{item.get('level')}] {item.get('description', '')}")
        run.bold = True
        set_cell_bg(header_cells[0], "4472C4")
        
        # Result row
        row = table.add_row().cells
        row[1].merge(row[2])
        row[0].paragraphs[0].add_run("Result:").bold = True
        row[1].paragraphs[0].add_run(str(item.get("result", ""))).bold = True
        
        # Info section
        if item.get("info"):
            row = table.add_row().cells
            row[0].merge(row[2])
            row[0].paragraphs[0].add_run("Info:").bold = True
            
            row = table.add_row().cells
            row[0].merge(row[2])
            row[0].paragraphs[0].add_run(item.get("info", ""))
        
        # Solution section
        if item.get("solution"):
            row = table.add_row().cells
            row[0].merge(row[2])
            row[0].paragraphs[0].add_run("Solution:").bold = True
            
            row = table.add_row().cells
            row[0].merge(row[2])
            row[0].paragraphs[0].add_run(item.get("solution", ""))
        
        # Impact section
        if item.get("impact"):
            row = table.add_row().cells
            row[0].merge(row[2])
            row[0].paragraphs[0].add_run("Impact:").bold = True
            
            row = table.add_row().cells
            row[0].merge(row[2])
            row[0].paragraphs[0].add_run(item.get("impact", ""))
        
        # Policy values
        row = table.add_row().cells
        row[1].merge(row[2])
        row[0].paragraphs[0].add_run("Policy Value:").bold = True
        row[1].paragraphs[0].add_run(item.get("policy_value", ""))
        
        row = table.add_row().cells
        row[1].merge(row[2])
        row[0].paragraphs[0].add_run("Configured Value:").bold = True
        row[1].paragraphs[0].add_run(item.get("actual_value", ""))
        
        document.add_paragraph("")
    
    document.save(output_path)


def build_details_excel(items: List[dict], output_path: str, host_ip: Optional[str] = None) -> None:
    """Build Windows 7 detailed report in Excel format."""
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Windows 7 Details"
    
    # Define styles
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Column widths
    column_widths = [12, 12, 35, 12, 18, 18, 18, 18, 15, 15]
    for i, width in enumerate(column_widths, 1):
        worksheet.column_dimensions[chr(64 + i)].width = width
    
    # Headers (10 columns)
    headers = [
        "IP Address",
        "Check ID",
        "Description",
        "Level",
        "Result",
        "Info",
        "Solution",
        "Impact",
        "Policy Value",
        "Configured Value"
    ]
    
    for col_num, header in enumerate(headers, 1):
        cell = worksheet.cell(row=1, column=col_num)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # Data rows
    for row_num, item in enumerate(items, 2):
        row_data = [
            host_ip or "",
            str(item.get("check_number", ""))[:15],
            str(item.get("description", "")),
            str(item.get("level", "")),
            str(item.get("result", "")),
            str(item.get("info", "")),
            str(item.get("solution", "")),
            str(item.get("impact", "")),
            str(item.get("policy_value", "")),
            str(item.get("actual_value", "")),
        ]
        
        for col_num, value in enumerate(row_data, 1):
            cell = worksheet.cell(row=row_num, column=col_num)
            cell.value = value
            cell.border = border
            cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
    
    worksheet.freeze_panes = "A2"
    workbook.save(output_path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Generate Windows 7 Compliance Reports (Summary and Details) in Excel and Word formats."
    )
    parser.add_argument("scan_path", help="Path to Nessus .nessus/.xml file")
    parser.add_argument(
        "--output",
        default="windows7_compliance",
        help="Output file prefix (default: windows7_compliance). Will create 4 files with suffixes.",
    )
    parser.add_argument(
        "--format",
        choices=["both", "word", "excel"],
        default="both",
        help="Output format: both (default), word, or excel",
    )
    parser.add_argument(
        "--type",
        choices=["both", "summary", "details"],
        default="both",
        help="Report type: both (default), summary, or details",
    )
    return parser


def main() -> None:
    args = build_parser().parse_args()
    
    print(f"🔍 Analyzing: {args.scan_path}")
    host_ip, items = extract_windows7_items(args.scan_path)
    
    if not items:
        print("❌ No Windows 7 compliance checks found in this file.")
        return
    
    l1, l2, total = count_l1_l2(items)
    print(f"✅ Found {total} Windows 7 checks (L1: {l1}, L2: {l2})")
    
    output_dir = Path(args.output).parent
    output_base = Path(args.output).stem
    
    if output_dir and output_dir != Path("."):
        output_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = Path(".")
    
    # Generate reports based on format and type
    if args.type in ["summary", "both"]:
        if args.format in ["word", "both"]:
            word_file = output_dir / f"{output_base}_summary.docx"
            build_summary_word(items, str(word_file), host_ip)
            print(f"✅ Summary Word: {word_file}")
        
        if args.format in ["excel", "both"]:
            excel_file = output_dir / f"{output_base}_summary.xlsx"
            build_summary_excel(items, str(excel_file), host_ip)
            print(f"✅ Summary Excel: {excel_file}")
    
    if args.type in ["details", "both"]:
        if args.format in ["word", "both"]:
            word_file = output_dir / f"{output_base}_details.docx"
            build_details_word(items, str(word_file), host_ip)
            print(f"✅ Details Word: {word_file}")
        
        if args.format in ["excel", "both"]:
            excel_file = output_dir / f"{output_base}_details.xlsx"
            build_details_excel(items, str(excel_file), host_ip)
            print(f"✅ Details Excel: {excel_file}")
    
    print(f"\n📊 Reports generated successfully!")


if __name__ == "__main__":
    main()
