"""Generate a single Word file with compliance tables for all Nessus files in a folder."""
from __future__ import annotations

import argparse
from pathlib import Path
import re
from typing import Dict, List, Optional, Tuple

from docx import Document

from parser import parse_nessus

CHECK_PATTERN = re.compile(r"^(?P<check>\S+)\s+\((?P<level>L[12])\)\s+(?P<desc>.+)$")


def parse_check_name(check_name: Optional[str]) -> Tuple[str, str, str]:
    if not check_name:
        return "", "", ""

    match = CHECK_PATTERN.match(check_name)
    if match:
        return (
            match.group("check"),
            match.group("level"),
            match.group("desc"),
        )

    return "", "", check_name


def build_compliance_rows(findings: List[Dict]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for finding in findings:
        if not finding.get("compliance"):
            continue
        check_number, level, description = parse_check_name(finding.get("compliance_check_name"))
        rows.append(
            {
                "No": str(len(rows) + 1).zfill(2),
                "L1/L2": level,
                "Check number": check_number,
                "Description": description,
                "Result": finding.get("compliance_result") or "Unknown",
            }
        )
    return rows


def add_table(document: Document, rows: List[Dict[str, str]], host_ip: Optional[str]) -> None:
    if host_ip:
        document.add_paragraph(f"Host IP: {host_ip}")
    if not rows:
        document.add_paragraph("No compliance results found.")
        return

    headers = ["No", "L1/L2", "Check number", "Description", "Result"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header

    for row in rows:
        row_cells = table.add_row().cells
        for idx, header in enumerate(headers):
            row_cells[idx].text = row.get(header, "")


def build_report(scan_folder: Path, output_path: Path) -> None:
    document = Document()
    document.add_heading("Compliance Summary Report", level=1)

    scan_paths = sorted(scan_folder.glob("*.nessus")) + sorted(scan_folder.glob("*.xml"))
    if not scan_paths:
        raise SystemExit(f"No .nessus/.xml files found in {scan_folder}")

    for scan_path in scan_paths:
        findings = parse_nessus(str(scan_path))
        rows = build_compliance_rows(findings)
        host_ip = next((finding.get("host") for finding in findings if finding.get("host")), None)
        add_table(document, rows, host_ip)
        document.add_page_break()

    document.save(output_path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Create one Word file with compliance tables for all scans in a folder.")
    parser.add_argument("scan_folder", help="Folder containing .nessus/.xml files")
    parser.add_argument(
        "--output",
        default="compliance_summary_report.docx",
        help="Output .docx path (default: compliance_summary_report.docx)",
    )
    return parser


def main() -> None:
    args = build_parser().parse_args()
    build_report(Path(args.scan_folder), Path(args.output))
    print(f"Wrote compliance summary report to {args.output}")


if __name__ == "__main__":
    main()
