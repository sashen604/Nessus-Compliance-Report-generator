"""Generate detailed compliance report tables in a Word document."""
from __future__ import annotations

import argparse
import re
import xml.etree.ElementTree as ET
from typing import Iterable, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

CHECK_PATTERN = re.compile(r"^(?P<check>\S+)\s+\((?P<level>L[12])\)\s+(?P<desc>.+)$")


def _strip_text(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    value = value.strip()
    return value if value else None


def _find_text(element: ET.Element, tag: str) -> Optional[str]:
    child = element.find(tag)
    if child is not None and child.text is not None:
        return _strip_text(child.text)

    for child in element:
        if child.tag.endswith(tag):
            if child.text is not None:
                return _strip_text(child.text)

    return None


def parse_check_name(check_name: Optional[str]) -> Tuple[str, str, str]:
    if not check_name:
        return "", "", ""

    match = CHECK_PATTERN.match(check_name)
    if match:
        return (
            match.group("check"),
            match.group("level"),
            match.group("desc"),
        )

    return "", "", check_name


def extract_items(scan_path: str) -> Tuple[Optional[str], List[dict]]:
    tree = ET.parse(scan_path)
    root = tree.getroot()

    host_ip: Optional[str] = None
    for report_host in root.iter("ReportHost"):
        host_ip = report_host.attrib.get("name") or _find_text(report_host, "host-ip")
        if host_ip:
            break

    items: List[dict] = []

    for report_item in root.iter("ReportItem"):
        compliance_text = _find_text(report_item, "compliance")
        compliance_result = _find_text(report_item, "compliance-result")
        if compliance_text and compliance_text.lower() != "true" and not compliance_result:
            continue
        if compliance_text is None and compliance_result is None:
            continue

        check_name = _find_text(report_item, "compliance-check-name")
        check_number, level, description = parse_check_name(check_name)

        info_text = _find_text(report_item, "compliance-info") or ""
        solution_text = _find_text(report_item, "compliance-solution") or ""
        impact_text = _find_text(report_item, "compliance-impact") or ""

        if not impact_text and solution_text:
            split_token = "Impact:"
            if split_token in solution_text:
                solution_part, impact_part = solution_text.split(split_token, 1)
                solution_text = solution_part.strip()
                impact_text = impact_part.strip()

        items.append(
            {
                "check_number": check_number,
                "level": level,
                "description": description,
                "result": compliance_result or "Unknown",
                "info": info_text,
                "solution": solution_text,
                "impact": impact_text,
                "policy_value": _find_text(report_item, "compliance-policy-value") or "",
                "actual_value": _find_text(report_item, "compliance-actual-value") or "",
            }
        )

    return host_ip, items


def _add_two_column_row(table, label: str, value: str, bold_label: bool = False, bold_value: bool = False) -> None:
    row_cells = table.add_row().cells
    row_cells[1].merge(row_cells[2])

    label_paragraph = row_cells[0].paragraphs[0]
    label_paragraph.clear()
    label_run = label_paragraph.add_run(label)
    label_run.bold = bold_label

    value_paragraph = row_cells[1].paragraphs[0]
    value_paragraph.clear()
    value_run = value_paragraph.add_run(value)
    value_run.bold = bold_value

def set_cell_bg(cell, color: str = "D9E1F2") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def set_column_widths(table) -> None:
    widths = [Inches(1.2), Inches(1.0), Inches(4.8)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width


def _add_full_width_row(table, text: str, bold: bool = False) -> None:
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[1]).merge(row_cells[2])
    paragraph = row_cells[0].paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = bold


def build_docx(items, output_path, host_ip=None):
    document = Document()
    if host_ip:
        p = document.add_paragraph(f"Host: {host_ip}")
        p.runs[0].bold = True
    document.add_heading("Compliance Detailed Report", level=1)
    for item in items:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        set_column_widths(table)
        # HEADER ROW
        header_cells = table.rows[0].cells
        header_values = [
            item.get("check_number") or "",
            f"({item.get('level')})" if item.get("level") else "",
            item.get("description") or "",
        ]
        for i, val in enumerate(header_values):
            p = header_cells[i].paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.bold = True
            set_cell_bg(header_cells[i], "B4C6E7")
        # RESULT ROW
        row = table.add_row().cells
        row[1].merge(row[2])
        r1 = row[0].paragraphs[0].add_run("Result")
        r1.bold = True
        r2 = row[1].paragraphs[0].add_run(item.get("result", ""))
        r2.bold = True
        # SECTION FUNCTION
        def add_section(title, content):
            # Title row
            r = table.add_row().cells
            r[0].merge(r[1]).merge(r[2])
            p = r[0].paragraphs[0]
            p.clear()
            run = p.add_run(title)
            run.bold = True
            # Content row
            r = table.add_row().cells
            r[0].merge(r[1]).merge(r[2])
            p = r[0].paragraphs[0]
            p.clear()
            p.add_run(content)
        add_section("Info", item.get("info", ""))
        add_section("Solution", item.get("solution", ""))
        add_section("Impact", item.get("impact", ""))
        # POLICY VALUE
        row = table.add_row().cells
        row[1].merge(row[2])
        row[0].paragraphs[0].add_run("Policy Value").bold = True
        row[1].paragraphs[0].add_run(item.get("policy_value", ""))
        # CONFIGURED VALUE
        row = table.add_row().cells
        row[1].merge(row[2])
        row[0].paragraphs[0].add_run("Configured Value").bold = True
        row[1].paragraphs[0].add_run(item.get("actual_value", ""))
        document.add_paragraph("")
    document.save(output_path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Create a detailed compliance report docx from Nessus XML.")
    parser.add_argument("scan_path", help="Path to Nessus .nessus/.xml file")
    parser.add_argument(
        "--output",
        default="compliance_detailed_report.docx",
        help="Output .docx path (default: compliance_detailed_report.docx)",
    )
    return parser


def main() -> None:
    args = build_parser().parse_args()
    host_ip, items = extract_items(args.scan_path)
    build_docx(items, args.output, host_ip=host_ip)
    print(f"Wrote {len(items)} detailed compliance tables to {args.output}")


if __name__ == "__main__":
    main()
