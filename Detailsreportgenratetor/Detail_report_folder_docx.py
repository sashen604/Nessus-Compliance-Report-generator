from __future__ import annotations
 
import argparse
import os
import re
import xml.etree.ElementTree as ET
from typing import List, Optional, Tuple
 
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
 
 
# =========================
# REGEX FOR CHECK PARSING
# =========================
CHECK_PATTERN = re.compile(r"^(?P<check>\S+)\s+\((?P<level>L[12])\)\s+(?P<desc>.+)$")
 
 
# =========================
# HELPER FUNCTIONS
# =========================
def _strip_text(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    value = value.strip()
    return value if value else None
 
 
def _find_text(element: ET.Element, tag: str) -> Optional[str]:
    child = element.find(tag)
    if child is not None and child.text is not None:
        return _strip_text(child.text)
 
    for child in element:
        if child.tag.endswith(tag):
            if child.text is not None:
                return _strip_text(child.text)
 
    return None
 
 
def parse_check_name(check_name: Optional[str]) -> Tuple[str, str, str]:
    if not check_name:
        return "", "", ""
 
    match = CHECK_PATTERN.match(check_name)
    if match:
        return (
            match.group("check"),
            match.group("level"),
            match.group("desc"),
        )
 
    return "", "", check_name
 
 
# =========================
# EXTRACT FROM NESSUS
# =========================
def extract_items(scan_path: str) -> Tuple[Optional[str], List[dict]]:
    tree = ET.parse(scan_path)
    root = tree.getroot()
 
    host_ip: Optional[str] = None
    for report_host in root.iter("ReportHost"):
        host_ip = report_host.attrib.get("name") or _find_text(report_host, "host-ip")
        if host_ip:
            break
 
    items: List[dict] = []
 
    for report_item in root.iter("ReportItem"):
        compliance_text = _find_text(report_item, "compliance")
        compliance_result = _find_text(report_item, "compliance-result")
 
        if compliance_text and compliance_text.lower() != "true" and not compliance_result:
            continue
        if compliance_text is None and compliance_result is None:
            continue
 
        check_name = _find_text(report_item, "compliance-check-name")
        check_number, level, description = parse_check_name(check_name)
 
        info_text = _find_text(report_item, "compliance-info") or ""
        solution_text = _find_text(report_item, "compliance-solution") or ""
        impact_text = _find_text(report_item, "compliance-impact") or ""
 
        if not impact_text and solution_text:
            if "Impact:" in solution_text:
                solution_part, impact_part = solution_text.split("Impact:", 1)
                solution_text = solution_part.strip()
                impact_text = impact_part.strip()
 
        items.append(
            {
                "check_number": check_number,
                "level": level,
                "description": description,
                "result": compliance_result or "Unknown",
                "info": info_text,
                "solution": solution_text,
                "impact": impact_text,
                "policy_value": _find_text(report_item, "compliance-policy-value") or "",
                "actual_value": _find_text(report_item, "compliance-actual-value") or "",
            }
        )
 
    return host_ip, items
 
 
# =========================
# WORD STYLING
# =========================
def set_cell_bg(cell, color="D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)
 
 
def set_column_widths(table):
    widths = [Inches(1.2), Inches(1.0), Inches(4.8)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width
 
 
# =========================
# BUILD SINGLE DOCUMENT
# =========================
def process_folder(folder_path: str, output_path: str):
    document = Document()
    document.add_heading("Combined Compliance Detailed Report", 0)
 
    total_items = 0
    total_files = 0
 
    for file in sorted(os.listdir(folder_path)):
        if file.endswith(".nessus") or file.endswith(".xml"):
            total_files += 1
            full_path = os.path.join(folder_path, file)
 
            print(f"Processing: {file}")
 
            host_ip, items = extract_items(full_path)
 
            # New section per scan
            document.add_page_break()
            document.add_heading(f"Scan File: {file}", level=1)
 
            if host_ip:
                p = document.add_paragraph(f"Host: {host_ip}")
                p.runs[0].bold = True
 
            for item in items:
                table = document.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                set_column_widths(table)
 
                # HEADER
                header_cells = table.rows[0].cells
                header_values = [
                    item["check_number"],
                    f"({item['level']})" if item["level"] else "",
                    item["description"],
                ]
 
                for i, val in enumerate(header_values):
                    p = header_cells[i].paragraphs[0]
                    p.clear()
                    run = p.add_run(val)
                    run.bold = True
                    set_cell_bg(header_cells[i], "B4C6E7")
 
                # RESULT
                row = table.add_row().cells
                row[1].merge(row[2])
                row[0].paragraphs[0].add_run("Result").bold = True
                row[1].paragraphs[0].add_run(item["result"]).bold = True
 
                def add_section(title, content):
                    r = table.add_row().cells
                    r[0].merge(r[1]).merge(r[2])
                    p = r[0].paragraphs[0]
                    p.clear()
                    p.add_run(title).bold = True
 
                    r = table.add_row().cells
                    r[0].merge(r[1]).merge(r[2])
                    p = r[0].paragraphs[0]
                    p.clear()
                    p.add_run(content)
 
                add_section("Info", item["info"])
                add_section("Solution", item["solution"])
                add_section("Impact", item["impact"])
 
                # POLICY VALUE
                row = table.add_row().cells
                row[1].merge(row[2])
                row[0].paragraphs[0].add_run("Policy Value").bold = True
                row[1].paragraphs[0].add_run(item["policy_value"])
 
                # CONFIGURED VALUE
                row = table.add_row().cells
                row[1].merge(row[2])
                row[0].paragraphs[0].add_run("Configured Value").bold = True
                row[1].paragraphs[0].add_run(item["actual_value"])
 
                document.add_paragraph("")
 
            total_items += len(items)
 
    document.save(output_path)
 
    print("\n==============================")
    print(f"✅ Files processed : {total_files}")
    print(f"📊 Total checks    : {total_items}")
    print(f"📄 Output file     : {output_path}")
    print("==============================")
 
 
# =========================
# MAIN
# =========================
def main():
    parser = argparse.ArgumentParser(
        description="Combine multiple Nessus compliance reports into one Word document"
    )
    parser.add_argument("folder_path", help="Folder containing .nessus files")
    parser.add_argument(
        "--output",
        default="combined_compliance_report.docx",
        help="Output Word file"
    )
 
    args = parser.parse_args()
 
    process_folder(args.folder_path, args.output)
 
 
if __name__ == "__main__":
    main()