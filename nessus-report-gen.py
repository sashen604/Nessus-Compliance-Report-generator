#!/usr/bin/env python3

from pathlib import Path
from datetime import datetime
from docx import Document

# SUMMARY TABLE
from compliance_table_docx import extract_rows, build_docx

# DETAILED PARSER + STYLE (YOUR GOOD FORMAT)
from details_report_docx import extract_items, set_cell_bg, set_column_widths

# EXCEL EXPORT
from details_report_excel import build_excel, build_excel_combined

# COUNT
from utils import count_l1_l2

# CLI SUMMARY
from parser import parse_nessus
from summarizer import (
    overall_summary,
    host_summary,
    compliance_summary,
    failed_compliance_checks,
)
from output import print_table


# =========================
# USER INPUT
# =========================
def ask_user():
    print("\n====== Nessus Report Generator ======\n")

    path = input("1. Enter .nessus file OR folder path: ").strip()

    print("\n2. Report Type:")
    print("   1 - Summary (Table)")
    print("   2 - Detailed (Professional DOCX)")
    print("   3 - Detailed (Excel)")
    report_type = input("Choose (1/2/3): ").strip()

    print("\n3. Output Mode:")
    print("   1 - Combined")
    print("   2 - Separate file per report")
    output_mode = input("Choose (1/2): ").strip()

    print("\n4. Level Filter:")
    print("   1 - L1 + L2")
    print("   2 - Only L1")
    level_filter = input("Choose (1/2): ").strip()

    output_path = input("\n5. Output folder (Enter = default): ").strip()
    output_name = input("6. Output file name (Enter = auto): ").strip()

    return path, report_type, output_mode, level_filter, output_path, output_name


# =========================
# DEFAULT NAME
# =========================
def generate_name(report_type):
    now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    if report_type == "1":
        return f"summary_{now}.docx"
    elif report_type == "3":
        return f"detailed_{now}.xlsx"
    else:
        return f"detailed_{now}.docx"


# =========================
# FILTER
# =========================
def filter_level(data, level_filter):
    if level_filter == "1":
        return data

    return [
        item for item in data
        if (item[1] if isinstance(item, tuple) else item.get("level")) == "L1"
    ]


# =========================
# CLI SUMMARY
# =========================
def run_summary(scan_path):
    findings = parse_nessus(scan_path)

    print("\n📊 ===== SUMMARY =====")

    overall = overall_summary(findings)
    rows = [{"Metric": k, "Count": v} for k, v in overall.items()]
    print_table("Overall Summary", rows)

    print_table("Host Summary", host_summary(findings))
    print_table("Compliance Summary", compliance_summary(findings))
    print_table("Failed Checks", failed_compliance_checks(findings))


# =========================
# ✅ PROFESSIONAL DETAILED REPORT (YOUR FORMAT)
# =========================
def build_detailed_pro_report(document, items, host_ip):
    if host_ip:
        p = document.add_paragraph(f"Host: {host_ip}")
        p.runs[0].bold = True

    l1, l2, total = count_l1_l2(items)
    document.add_paragraph(f"L1: {l1} | L2: {l2} | Total: {total}")

    for item in items:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        set_column_widths(table)

        # HEADER
        header_cells = table.rows[0].cells
        values = [
            item["check_number"],
            f"({item['level']})",
            item["description"],
        ]

        for i, val in enumerate(values):
            p = header_cells[i].paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.bold = True
            set_cell_bg(header_cells[i], "B4C6E7")

        # RESULT
        row = table.add_row().cells
        row[1].merge(row[2])
        row[0].paragraphs[0].add_run("Result").bold = True
        row[1].paragraphs[0].add_run(item["result"]).bold = True

        # SECTIONS
        def add_section(title, content):
            r = table.add_row().cells
            r[0].merge(r[1]).merge(r[2])
            r[0].paragraphs[0].add_run(title).bold = True

            r = table.add_row().cells
            r[0].merge(r[1]).merge(r[2])
            r[0].paragraphs[0].add_run(content)

        add_section("Info", item["info"])
        add_section("Solution", item["solution"])
        add_section("Impact", item["impact"])

        # POLICY
        row = table.add_row().cells
        row[1].merge(row[2])
        row[0].paragraphs[0].add_run("Policy Value").bold = True
        row[1].paragraphs[0].add_run(item["policy_value"])

        # CONFIGURED
        row = table.add_row().cells
        row[1].merge(row[2])
        row[0].paragraphs[0].add_run("Configured Value").bold = True
        row[1].paragraphs[0].add_run(item["actual_value"])

        document.add_paragraph("")


# =========================
# PROCESS FILE
# =========================
def process_file(file_path, report_type, level_filter, output_file):
    print(f"\n📂 Processing: {file_path}")

    run_summary(file_path)

    if report_type == "1":
        host_ip, rows = extract_rows(file_path)
        rows = filter_level(rows, level_filter)
        build_docx(rows, output_file, host_ip)

    elif report_type == "3":
        # Excel report
        host_ip, items = extract_items(file_path)
        items = filter_level(items, level_filter)
        build_excel(items, output_file, host_ip)

    else:
        # Detailed DOCX report
        document = Document()
        document.add_heading("Compliance Detailed Report", 0)

        host_ip, items = extract_items(file_path)
        items = filter_level(items, level_filter)

        build_detailed_pro_report(document, items, host_ip)

        document.save(output_file)


# =========================
# MAIN
# =========================
def main():
    path, report_type, output_mode, level_filter, output_path, output_name = ask_user()

    path = Path(path)

    output_dir = Path(output_path) if output_path else Path("output")
    output_dir.mkdir(parents=True, exist_ok=True)

    if not output_name:
        output_name = generate_name(report_type)

    # Ensure correct file extension based on report type
    if report_type == "3":
        if not output_name.endswith(".xlsx"):
            output_name = output_name.replace(".docx", "") + ".xlsx" if ".docx" in output_name else output_name + ".xlsx"
    else:
        if not output_name.endswith(".docx"):
            output_name = output_name.replace(".xlsx", "") + ".docx" if ".xlsx" in output_name else output_name + ".docx"

    # SINGLE FILE
    if path.is_file():
        output_file = output_dir / output_name
        process_file(str(path), report_type, level_filter, str(output_file))
        print(f"\n📄 Saved: {output_file}")

    # FOLDER
    else:
        files = list(path.glob("*.nessus"))

        if not files:
            print("❌ No files found")
            return

        if output_mode == "2":
            for f in files:
                file_out = output_dir / (f.stem + "_" + output_name)
                process_file(str(f), report_type, level_filter, str(file_out))

        else:
            # For combined mode with Excel, create one combined file
            if report_type == "3":
                print("📊 Creating combined Excel file...")
                all_files_data = []
                
                for f in files:
                    print(f"   📂 Processing: {f.name}")
                    host_ip, items = extract_items(str(f))
                    items = filter_level(items, level_filter)
                    all_files_data.append((f.name, host_ip, items))
                
                output_file = output_dir / output_name
                build_excel_combined(all_files_data, str(output_file))
                print(f"\n✅ Combined Excel saved: {output_file}")
            else:
                document = Document()
                document.add_heading("Combined Nessus Report", 0)

                for f in files:
                    document.add_page_break()
                    document.add_heading(f"File: {f.name}", level=1)

                    if report_type == "1":
                        host_ip, rows = extract_rows(str(f))
                        rows = filter_level(rows, level_filter)

                        table = document.add_table(rows=1, cols=5)
                        table.style = "Table Grid"

                        headers = ["No", "L1/L2", "Check number", "Description", "Status"]
                        for i, h in enumerate(headers):
                            table.rows[0].cells[i].text = h

                        for idx, row in enumerate(rows, start=1):
                            cells = table.add_row().cells
                            cells[0].text = f"{idx:02d}"
                            cells[1].text = row[1]
                            cells[2].text = row[0]
                            cells[3].text = row[2]
                            cells[4].text = row[3]

                    else:
                        host_ip, items = extract_items(str(f))
                        items = filter_level(items, level_filter)
                        build_detailed_pro_report(document, items, host_ip)

                output_file = output_dir / output_name
                document.save(output_file)

                print(f"\n📄 Combined saved: {output_file}")


# =========================
# ENTRY
# =========================
if __name__ == "__main__":
    main()