"""Generate a Word table from Nessus compliance checks."""
from __future__ import annotations

import argparse
import re
import xml.etree.ElementTree as ET
from utils import count_l1_l2
from typing import List, Optional, Tuple

from docx import Document


CHECK_PATTERN = re.compile(r"^(?P<check>\S+)\s+\((?P<level>L[12])\)\s+(?P<desc>.+)$")


def _strip_text(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    value = value.strip()
    return value if value else None


def _find_text(element: ET.Element, tag: str) -> Optional[str]:
    child = element.find(tag)
    if child is not None and child.text is not None:
        return _strip_text(child.text)

    for child in element:
        if child.tag.endswith(tag):
            if child.text is not None:
                return _strip_text(child.text)

    return None


def parse_check_name(check_name: Optional[str]) -> Tuple[str, str, str]:
    if not check_name:
        return "", "", ""

    match = CHECK_PATTERN.match(check_name)
    if match:
        return (
            match.group("check"),
            match.group("level"),
            match.group("desc"),
        )

    return "", "", check_name


def extract_rows(scan_path: str) -> Tuple[Optional[str], List[Tuple[str, str, str, str]]]:
    tree = ET.parse(scan_path)
    root = tree.getroot()

    rows: List[Tuple[str, str, str, str]] = []
    host_ip: Optional[str] = None

    for report_host in root.iter("ReportHost"):
        host_ip = report_host.attrib.get("name") or _find_text(report_host, "host-ip")
        if host_ip:
            break

    for report_item in root.iter("ReportItem"):
        compliance_text = _find_text(report_item, "compliance")
        compliance_result = _find_text(report_item, "compliance-result")
        if compliance_text and compliance_text.lower() != "true" and not compliance_result:
            continue
        if compliance_text is None and compliance_result is None:
            continue

        check_name = _find_text(report_item, "compliance-check-name") or _find_text(report_item, "cm:compliance-check-name")
        check_number, level, description = parse_check_name(check_name)
        
        # If level wasn't found from check_name format, try to get it from cm:compliance-benchmark-profile
        if not level:
            level = _find_text(report_item, "cm:compliance-benchmark-profile") or _find_text(report_item, "compliance-benchmark-profile") or ""
            # If still no level but we have description, use description as is
            if not level and not description:
                description = check_name or ""
        elif not description:
            # If we got level from parse_check_name but not description, use check_name
            description = check_name or ""
            
        status = compliance_result or _find_text(report_item, "cm:compliance-result") or _find_text(report_item, "compliance-result") or _find_text(report_item, "cm:compliance-actual-value") or "Unknown"
        
        check_number = check_number or _find_text(report_item, "cm:compliance-check-id") or _find_text(report_item, "compliance-check-id") or ""

        if not check_number and not level and not description:
            continue

        rows.append((check_number, level, description, status))

    return host_ip, rows


def build_docx(rows, output_path, host_ip=None):
    document = Document()
    document.add_heading("Compliance Checks", level=1)

    if host_ip:
        document.add_paragraph(f"Host IP: {host_ip}")

    # ✅ ADD L1/L2 COUNTS
    l1, l2, total = count_l1_l2(rows)

    document.add_paragraph(f"L1 Checks : {l1}")
    document.add_paragraph(f"L2 Checks : {l2}")
    document.add_paragraph(f"Total     : {total}")

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = ["No", "L1/L2", "Check number", "Description", "Status"]

    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for idx, (check_number, level, description, status) in enumerate(rows, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = f"{idx:02d}"
        row_cells[1].text = level
        row_cells[2].text = check_number
        row_cells[3].text = description
        row_cells[4].text = status

    document.save(output_path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Create a compliance table docx from Nessus XML.")
    parser.add_argument("scan_path", help="Path to Nessus .nessus/.xml file")
    parser.add_argument(
        "--output",
        default="compliance_table.docx",
        help="Output .docx path (default: compliance_table.docx)",
    )
    return parser


def main() -> None:
    args = build_parser().parse_args()
    host_ip, rows = extract_rows(args.scan_path)
    build_docx(rows, args.output, host_ip=host_ip)
    print(f"Wrote {len(rows)} compliance rows to {args.output}")


if __name__ == "__main__":
    main()
